import os
import logging
from concurrent.futures import ThreadPoolExecutor
import threading
import streamlit as st
from utils.query_engine import QueryEngine

logger = logging.getLogger(__name__)


def handle_file_upload(uploaded_file, session_id, db):
    """Handle file upload with optimized 1MB chunking."""
    try:
        UPLOAD_DIR = "data/uploads"
        session_upload_dir = os.path.join(UPLOAD_DIR, session_id)
        os.makedirs(session_upload_dir, exist_ok=True)
        file_path = os.path.join(session_upload_dir, uploaded_file.name)
        if db.add_file(session_id, file_path, uploaded_file.name):
            CHUNK_SIZE = 1024 * 1024  # 1MB
            with open(file_path, "wb") as f:
                while True:
                    chunk = uploaded_file.read(CHUNK_SIZE)
                    if not chunk:
                        break
                    f.write(chunk)
            return True
        return False
    except Exception:
        return False


def process_uploads(files, session_id, db):
    """Process multiple file uploads concurrently."""
    results = []
    with ThreadPoolExecutor(max_workers=min(len(files), 4)) as executor:
        futures = {
            executor.submit(handle_file_upload, f, session_id, db): f for f in files
        }
        for future in futures:
            results.append(future.result())
    # Trigger background indexing so the user can interact immediately.
    query_engine = st.session_state["query_engines"].get(session_id)
    if query_engine:
        threading.Thread(
            target=query_engine.index_manager.build_index, daemon=True
        ).start()
    return results


def rerun_assistant_message(session_id, assistant_timestamp, user_question, db):
    """
    Rerun the assistant response for a given user question.
    Delete the existing assistant message (identified by its timestamp)
    and add a new response from the query engine.
    """
    delete_chat_message(session_id, assistant_timestamp, db)
    # Reinitialize the query engine to force a fresh response.
    st.session_state["query_engines"][session_id] = QueryEngine(
        os.getenv("GROQ_API_KEY"), session_id=session_id
    )
    query_engine = st.session_state["query_engines"].get(session_id)
    if query_engine:
        new_response = query_engine.query(user_question, conversation_history=[])
        db.add_message(session_id, "assistant", new_response)


def delete_chat_message(session_id, timestamp, db):
    """Delete a single chat message identified by its timestamp."""
    try:
        cur = db.conn.cursor()
        cur.execute(
            "DELETE FROM messages WHERE session_id = ? AND timestamp = ?",
            (session_id, timestamp),
        )
        db.conn.commit()
        cur.close()
        return True
    except Exception:
        return False


def delete_chat_message_pair(session_id, user_timestamp, assistant_timestamp, db):
    """Delete a pair of chat messages (user and assistant) identified by their timestamps."""
    success_user = delete_chat_message(session_id, user_timestamp, db)
    success_assistant = delete_chat_message(session_id, assistant_timestamp, db)
    return success_user and success_assistant


def save_note_and_get_path(
    user_question: str, assistant_response: str, note_title: str, file_type: str, db
) -> str:
    """Save a note with proper formatting."""
    try:
        NOTES_DIR = os.path.abspath("data/notes")
        os.makedirs(NOTES_DIR, exist_ok=True)

        # Clean up title and ensure valid filename
        safe_title = "".join(
            c for c in note_title if c.isalnum() or c in (" ", "-", "_")
        ).strip()
        filename = f"{safe_title}.{file_type}"  # Remove timestamp from filename
        file_path = os.path.join(NOTES_DIR, filename)

        # Content with proper formatting
        formatted_content = f"""# {user_question}
{assistant_response}"""

        try:
            if file_type == "txt":
                with open(file_path, "w", encoding="utf-8", newline="\n") as f:
                    f.write(formatted_content)

            elif file_type == "pdf":
                try:
                    from fpdf import FPDF

                    pdf = FPDF()
                    pdf.add_page()

                    # Title (User Question)
                    pdf.set_font("Arial", "B", size=16)
                    # Split long titles into multiple lines
                    pdf.multi_cell(0, 10, user_question)
                    pdf.ln(10)

                    # Content (Assistant Response)
                    pdf.set_font("Arial", "", size=12)
                    pdf.multi_cell(0, 10, assistant_response)

                    pdf.output(file_path)
                except ImportError:
                    st.error(
                        "PDF creation requires fpdf2. Please install it with: pip install fpdf2"
                    )
                    return None

            elif file_type == "docx":
                try:
                    from docx import Document
                    from docx.shared import Pt
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                    doc = Document()

                    # Title (User Question)
                    heading = doc.add_heading(level=1)
                    heading_run = heading.add_run(user_question)
                    heading_run.font.size = Pt(16)
                    heading_run.font.bold = True
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    # Add some space
                    doc.add_paragraph()

                    # Content (Assistant Response)
                    content_para = doc.add_paragraph()
                    content_run = content_para.add_run(assistant_response)
                    content_run.font.size = Pt(12)

                    doc.save(file_path)
                except ImportError:
                    st.error(
                        "DOCX creation requires python-docx. Please install it with: pip install python-docx"
                    )
                    return None

            # Add note to database
            if db.add_note(
                title=safe_title,  # Use clean title without timestamp
                content=formatted_content,
                file_path=file_path,
                file_type=file_type,
                conversation_id=st.session_state.get("selected_session_id"),
            ):
                return file_path
            else:
                if os.path.exists(file_path):
                    os.remove(file_path)
                return None

        except Exception as e:
            logger.error(f"Error saving note file: {e}")
            if os.path.exists(file_path):
                os.remove(file_path)
            return None

    except Exception as e:
        logger.error(f"Error in save_note_and_get_path: {e}")
        return None
