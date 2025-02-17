import streamlit as st
import os
import time
from typing import Optional, Tuple
import pymupdf  # Correct import for PyMuPDF
from docx import Document
import logging
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileReader:
    """Class to handle different file types reading operations."""

    @staticmethod
    def read_txt(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Read content from a text file."""
        try:
            with open(file_path, "rb") as f:
                content = f.read().decode("utf-8", errors="replace")
            return content, None
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return None, f"Error reading text file: {str(e)}"

    @staticmethod
    def read_pdf(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Read content from a PDF file with proper formatting."""
        try:
            pdf_content = []
            doc = pymupdf.open(file_path)

            # Extract a title from the first page (if available)
            first_page = doc[0]
            title = first_page.get_text(
                "text", clip=(50, 0, first_page.rect.width - 50, 100)
            ).strip()

            for page in doc:
                content = page.get_text().strip()
                # Remove title from first page to avoid duplication
                if page.number == 0 and title:
                    content = content.replace(title, "", 1)
                pdf_content.append(content)
            doc.close()

            formatted_content = f"# {title}\n\n" + "\n".join(pdf_content)

            return formatted_content, None
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return None, f"Error reading PDF file: {str(e)}"

    @staticmethod
    def read_docx(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Read content from a DOCX file with proper formatting."""
        try:
            doc = Document(file_path)
            title = doc.paragraphs[0].text if doc.paragraphs else ""
            content = "\n".join(p.text for p in doc.paragraphs[1:])
            formatted_content = f"# {title}\n\n{content}"
            return formatted_content, None
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return None, f"Error reading DOCX file: {str(e)}"


def get_file_content(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Get content from a file based on its extension."""
    file_extension = os.path.splitext(file_path)[1].lower()
    file_handlers = {
        ".txt": FileReader.read_txt,
        ".pdf": FileReader.read_pdf,
        ".docx": FileReader.read_docx,
    }
    handler = file_handlers.get(file_extension)
    if handler:
        return handler(file_path)
    return None, f"Unsupported file type: {file_extension}"


def render_notes():
    """Render the Notes tab with improved note management."""
    st.header("📝 Saved Notes")
    db = st.session_state.get("db")
    if not db:
        st.error("Database connection not available")
        return

    notes = db.get_notes()
    if not notes:
        st.info(
            "No notes saved yet. Use the 'Note' button under an AI response to save one."
        )
        return

    # Ensure state keys exist
    st.session_state.note_rename = st.session_state.get("note_rename", {})
    st.session_state.pending_note_deletion = st.session_state.get(
        "pending_note_deletion", None
    )

    for title, content, file_path, file_type, created_at, conversation_id in notes:
        if not os.path.exists(file_path):
            db.delete_note(file_path)
            continue

        with st.expander(f"📝 {title}.{file_type}", expanded=False):
            created_date = datetime.fromisoformat(created_at).strftime(
                "%B %d, %Y at %H:%M"
            )
            st.caption(f"Created on {created_date}")

            with st.spinner("Loading note content..."):
                file_content, error = get_file_content(file_path)
            if error:
                st.error(error)
            elif file_content:
                try:
                    parts = file_content.split("\n\n", 1)
                    if len(parts) == 2:
                        title_line, body = parts
                        st.markdown(
                            f"<h2>{title_line.replace('#', '').strip()}</h2>",
                            unsafe_allow_html=True,
                        )
                        st.markdown(body)
                    else:
                        st.markdown(file_content)
                except Exception as e:
                    logger.error(f"Error formatting note content: {e}")
                    st.markdown(file_content)

            col1, col2, col3 = st.columns(3)
            # Rename functionality
            with col1:
                if file_path in st.session_state.note_rename:
                    new_title = st.text_input(
                        "New title", value=title, key=f"new_title_{file_path}"
                    )
                    if st.button("Save", key=f"save_rename_{file_path}"):
                        try:
                            # Get new file path
                            new_file_path = os.path.join(
                                os.path.dirname(file_path), f"{new_title}.{file_type}"
                            )

                            # First update the database
                            if db.update_note_title(file_path, new_title):
                                # Then rename the file
                                os.rename(file_path, new_file_path)
                                # Update the file path in database
                                db.update_note_path(file_path, new_file_path)
                                st.success("Title updated!")
                                del st.session_state.note_rename[file_path]
                                time.sleep(0.5)  # Brief pause to show success message
                                st.rerun()
                            else:
                                st.error("Failed to update note in database")
                        except Exception as e:
                            logger.error(f"Error renaming note: {e}")
                            st.error("Failed to rename note")

                    if st.button("Cancel", key=f"cancel_rename_{file_path}"):
                        del st.session_state.note_rename[file_path]
                        st.rerun()
                else:
                    if st.button("✏️ Rename", key=f"rename_{file_path}"):
                        st.session_state.note_rename[file_path] = True
                        st.rerun()

            # Download functionality
            with col2:
                try:
                    with open(file_path, "rb") as f:
                        st.download_button(
                            label="💾 Download",
                            data=f.read(),
                            file_name=os.path.basename(file_path),
                            key=f"download_{file_path}",
                        )
                except Exception as e:
                    st.error(f"Error loading file: {e}")

            # Delete functionality with confirmation
            with col3:
                if st.session_state.pending_note_deletion == file_path:
                    confirm_col, cancel_col = st.columns(2)
                    with confirm_col:
                        if st.button(
                            "✓", key=f"confirm_delete_{file_path}", help="Confirm"
                        ):
                            try:
                                if os.path.exists(file_path):
                                    os.remove(file_path)
                                if db.delete_note(file_path):
                                    st.success("Note deleted!")
                                    st.session_state.pending_note_deletion = None
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    st.error("Failed to delete note from database")
                            except Exception as e:
                                logger.error(f"Error deleting note: {e}")
                                st.error(f"Failed to delete note: {str(e)}")
                    with cancel_col:
                        if st.button(
                            "×", key=f"cancel_delete_{file_path}", help="Cancel"
                        ):
                            st.session_state.pending_note_deletion = None
                            st.rerun()
                else:
                    if st.button("🗑️ Delete", key=f"delete_{file_path}"):
                        st.session_state.pending_note_deletion = file_path
                        st.rerun()
